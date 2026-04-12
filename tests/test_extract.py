"""
Tests for the document extraction module (scripts/core/extract.py).
"""

from pathlib import Path

import pytest

from scripts.core.extract import (
    extract_file,
    get_supported_extensions,
    create_extract_transcript,
    UnsupportedFormatError,
)


def make_file(path: Path, content: str = "test") -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")
    return path


class TestRegistry:
    """Test the extractor registry."""

    def test_supported_extensions_includes_docx(self):
        exts = get_supported_extensions()
        assert ".docx" in exts

    def test_unsupported_extension_raises(self, tmp_path):
        path = make_file(tmp_path / "file.xyz", "content")
        with pytest.raises(UnsupportedFormatError):
            extract_file(path)


class TestDocxExtractor:
    """Test DOCX text extraction."""

    def _make_docx(self, path, paragraphs):
        from docx import Document
        doc = Document()
        for para in paragraphs:
            doc.add_paragraph(para)
        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(path))
        return path

    def test_extract_docx_text(self, tmp_path):
        path = self._make_docx(
            tmp_path / "letter.docx",
            ["Dear Alice,", "We drove to Springfield last week.", "Love, Bob"]
        )
        text, metadata = extract_file(path)
        assert "Dear Alice" in text
        assert "Springfield" in text
        assert "Love, Bob" in text

    def test_extract_docx_metadata(self, tmp_path):
        path = self._make_docx(tmp_path / "doc.docx", ["Hello world"])
        text, metadata = extract_file(path)
        assert metadata["format"] == "docx"
        assert metadata["word_count"] == 2

    def test_extract_docx_preserves_paragraphs(self, tmp_path):
        path = self._make_docx(tmp_path / "doc.docx", ["Paragraph one.", "Paragraph two."])
        text, metadata = extract_file(path)
        assert "Paragraph one." in text
        assert "Paragraph two." in text
        assert text.index("Paragraph one.") < text.index("Paragraph two.")

    def test_extract_docx_with_table(self, tmp_path):
        from docx import Document
        doc = Document()
        doc.add_paragraph("Before table")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Value"
        table.cell(1, 0).text = "Alice"
        table.cell(1, 1).text = "100"
        doc.add_paragraph("After table")
        path = tmp_path / "table.docx"
        doc.save(str(path))
        text, metadata = extract_file(path)
        assert "Name" in text
        assert "Alice" in text
        assert "100" in text


class TestTranscriptGeneration:
    """Test .transcript.md creation from extracted text."""

    def test_creates_transcript_file(self, tmp_path):
        dest = tmp_path / "archive"
        dest.mkdir()
        source = make_file(dest / "Letters" / "letter.docx", "dummy")
        md_path = create_extract_transcript(
            source, "Dear Alice, hello.", {"word_count": 3, "format": "docx"}, dest
        )
        assert md_path.exists()
        assert md_path.name == "letter.transcript.md"

    def test_transcript_has_frontmatter(self, tmp_path):
        dest = tmp_path / "archive"
        dest.mkdir()
        source = make_file(dest / "doc.docx", "dummy")
        md_path = create_extract_transcript(
            source, "Hello world.", {"word_count": 2, "format": "docx"}, dest
        )
        content = md_path.read_text(encoding="utf-8")
        assert "source_file: doc.docx" in content
        assert "transcription_confidence: high" in content
        assert "transcription_method: extract (docx)" in content
        assert "word_count: 2" in content

    def test_transcript_has_body(self, tmp_path):
        dest = tmp_path / "archive"
        dest.mkdir()
        source = make_file(dest / "doc.docx", "dummy")
        md_path = create_extract_transcript(
            source, "The full body text.", {"word_count": 4, "format": "docx"}, dest
        )
        content = md_path.read_text(encoding="utf-8")
        assert "The full body text." in content


class TestDocExtractor:
    """Test DOC (binary Word) text extraction."""

    def test_doc_extension_registered(self):
        exts = get_supported_extensions()
        assert ".doc" in exts

    def test_extract_invalid_doc_returns_empty(self, tmp_path):
        """Invalid DOC file should return empty text, not crash."""
        path = tmp_path / "bad.doc"
        path.write_bytes(b"\x00" * 100)
        text, metadata = extract_file(path)
        assert metadata["format"] == "doc"
        assert isinstance(text, str)


class TestXlsxExtractor:
    """Test XLSX text extraction."""

    def _make_xlsx(self, path, sheets):
        from openpyxl import Workbook
        wb = Workbook()
        first = True
        for name, rows in sheets.items():
            if first:
                ws = wb.active
                ws.title = name
                first = False
            else:
                ws = wb.create_sheet(name)
            for row in rows:
                ws.append(row)
        path.parent.mkdir(parents=True, exist_ok=True)
        wb.save(str(path))
        return path

    def test_extract_xlsx_text(self, tmp_path):
        path = self._make_xlsx(tmp_path / "data.xlsx", {
            "People": [["Name", "Age"], ["Alice", 30], ["Bob", 25]]
        })
        text, metadata = extract_file(path)
        assert "Alice" in text
        assert "Bob" in text
        assert "30" in text

    def test_extract_xlsx_metadata(self, tmp_path):
        path = self._make_xlsx(tmp_path / "data.xlsx", {
            "Sheet1": [["A", "B"], ["C", "D"]]
        })
        text, metadata = extract_file(path)
        assert metadata["format"] == "xlsx"
        assert metadata["sheet_count"] == 1
        assert metadata["word_count"] > 0

    def test_extract_xlsx_multiple_sheets(self, tmp_path):
        path = self._make_xlsx(tmp_path / "multi.xlsx", {
            "People": [["Name"], ["Alice"]],
            "Places": [["City"], ["Springfield"]],
        })
        text, metadata = extract_file(path)
        assert "Alice" in text
        assert "Springfield" in text
        assert metadata["sheet_count"] == 2
        assert "People" in text
        assert "Places" in text

    def test_extract_xlsx_empty_cells_skipped(self, tmp_path):
        path = self._make_xlsx(tmp_path / "sparse.xlsx", {
            "Sheet1": [["A", None, "B"], [None, None, None], ["C", "D", None]]
        })
        text, metadata = extract_file(path)
        assert "A" in text
        assert "B" in text
        assert "C" in text


class TestXlsExtractor:
    """Test XLS (legacy binary) text extraction."""

    def _make_xls(self, path, sheets):
        try:
            import xlwt
        except ImportError:
            pytest.skip("xlwt not installed")
        wb = xlwt.Workbook()
        for name, rows in sheets.items():
            ws = wb.add_sheet(name)
            for r, row in enumerate(rows):
                for c, val in enumerate(row):
                    if val is not None:
                        ws.write(r, c, val)
        path.parent.mkdir(parents=True, exist_ok=True)
        wb.save(str(path))
        return path

    def test_xls_extension_registered(self):
        exts = get_supported_extensions()
        assert ".xls" in exts

    def test_extract_xls_text(self, tmp_path):
        path = self._make_xls(tmp_path / "data.xls", {
            "People": [["Name", "Age"], ["Alice", 30]]
        })
        text, metadata = extract_file(path)
        assert "Alice" in text
        assert metadata["format"] == "xls"
        assert metadata["sheet_count"] == 1

    def test_extract_xls_multiple_sheets(self, tmp_path):
        path = self._make_xls(tmp_path / "multi.xls", {
            "Names": [["Alice"]],
            "Cities": [["Springfield"]],
        })
        text, metadata = extract_file(path)
        assert "Alice" in text
        assert "Springfield" in text
        assert metadata["sheet_count"] == 2
