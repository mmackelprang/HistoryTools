"""
Document text extraction for the Family Archive.

Pluggable registry of extractors that map file extensions to extraction
functions. Each extractor takes a file path and returns (text, metadata).
Adding a new format is one decorated function.
"""

import os
from pathlib import Path
from datetime import datetime

# ── Registry ──────────────────────────────────────────────────────────────

EXTRACTORS = {}


class UnsupportedFormatError(ValueError):
    """Raised when no extractor is registered for a file extension."""
    pass


def register(ext):
    """Decorator to register an extractor for a file extension."""
    def decorator(fn):
        EXTRACTORS[ext.lower()] = fn
        return fn
    return decorator


def extract_file(path):
    """Extract text and metadata from a supported file.

    Args:
        path: Path to the file.

    Returns:
        Tuple of (text, metadata) where text is the extracted content
        and metadata is a dict with keys like word_count, format.

    Raises:
        UnsupportedFormatError if no extractor registered for the extension.
    """
    path = Path(path)
    ext = path.suffix.lower()
    if ext not in EXTRACTORS:
        raise UnsupportedFormatError(f"No extractor for {ext}")
    return EXTRACTORS[ext](path)


def get_supported_extensions():
    """Return set of all registered file extensions."""
    return set(EXTRACTORS.keys())


# ── Transcript generation ─────────────────────────────────────────────────


def create_extract_transcript(file_path, text, metadata, dest_root=None):
    """Write a .transcript.md file for an extracted document.

    Args:
        file_path: Path to the source document.
        text: Extracted text body.
        metadata: Dict with word_count, format, and optional fields.
        dest_root: Unused (kept for API compatibility with pipeline callers).

    Returns:
        Path to the created .transcript.md file.
    """
    file_path = Path(file_path)
    md_path = file_path.with_suffix(".transcript.md")
    fmt = metadata.get("format", "unknown")
    word_count = metadata.get("word_count", len(text.split()))
    today = datetime.now().strftime("%Y-%m-%d")

    content = f"""---
source_file: {file_path.name}
transcription_date: {today}
transcription_confidence: high
transcription_method: extract ({fmt})
word_count: {word_count}
---

{text}
"""
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(content)
    return md_path


# ── Extractors ────────────────────────────────────────────────────────────


@register(".docx")
def extract_docx(path):
    """Extract text from a DOCX file using python-docx."""
    from docx import Document

    doc = Document(str(path))
    parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    text = "\n\n".join(parts)
    word_count = len(text.split())
    return text, {"word_count": word_count, "format": "docx"}


@register(".doc")
def extract_doc(path):
    """Extract text from a DOC (binary Word) file using olefile.

    DOC files store text in the WordDocument stream as a sequence of
    bytes. This extracts readable ASCII/UTF-8 text. Complex formatting
    and embedded objects are skipped.
    """
    import olefile

    text = ""
    try:
        if not olefile.isOleFile(str(path)):
            return "", {"word_count": 0, "format": "doc"}

        ole = olefile.OleFileIO(str(path))
        if ole.exists("WordDocument"):
            data = ole.openstream("WordDocument").read()
            text = data.decode("latin-1", errors="replace")
            text = "".join(c if c.isprintable() or c in "\n\r\t" else " " for c in text)
            lines = [line.strip() for line in text.split("\n") if line.strip()]
            text = "\n\n".join(lines)
        ole.close()
    except Exception:
        text = ""

    word_count = len(text.split()) if text else 0
    return text, {"word_count": word_count, "format": "doc"}


@register(".xlsx")
def extract_xlsx(path):
    """Extract text from an XLSX file using openpyxl."""
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True, data_only=True)
    parts = []
    sheet_count = len(wb.sheetnames)

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        if sheet_count > 1:
            parts.append(f"## Sheet: {sheet_name}")

        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                parts.append("\t".join(cells))

    wb.close()
    text = "\n\n".join(parts)
    word_count = len(text.split())
    return text, {"word_count": word_count, "format": "xlsx", "sheet_count": sheet_count}


@register(".xls")
def extract_xls(path):
    """Extract text from an XLS (legacy binary) file using xlrd."""
    import xlrd

    wb = xlrd.open_workbook(str(path))
    parts = []
    sheet_count = wb.nsheets

    for sheet_idx in range(sheet_count):
        ws = wb.sheet_by_index(sheet_idx)
        if sheet_count > 1:
            parts.append(f"## Sheet: {ws.name}")

        for row_idx in range(ws.nrows):
            cells = []
            for col_idx in range(ws.ncols):
                val = ws.cell_value(row_idx, col_idx)
                if val != "":
                    if isinstance(val, float) and val == int(val):
                        cells.append(str(int(val)))
                    else:
                        cells.append(str(val))
            if cells:
                parts.append("\t".join(cells))

    text = "\n\n".join(parts)
    word_count = len(text.split())
    return text, {"word_count": word_count, "format": "xls", "sheet_count": sheet_count}
